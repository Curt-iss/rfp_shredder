#!/usr/bin/env python3

# Imports --------------------------------------

import docx
from io import StringIO
from pathlib import Path
import sys

'''
doc = docx.Document('test.docx')
for block in iter_block_items(doc):
    print(block.text)
'''

# How To Debug Crap in VSCode:
#
#
# 1. Somewhere in the file you want to debug, find a line you want to "stop" on
#       (i.e. set a breakpoint), at the time of writing, scroll down to line 94
#       (yu could pick any line though, doesn't matter)
#       If you mouse over the line number on the left side of the editor, you'll
#       see a red dot. Click the dot to set a breakpoint.
# 2. Find the Ladybug with a play button on the the Left side of your window
#       Click on it
# 3. You should see a button that says "Run an Debug Python", click it or press F5
# 4. You'll see a popup in the middle of the screen, you want the first "Python File (Debug current file)"
# 5. Give it a sec and the editor should jump to line you set the breakpoint on. You can move forward
#       line by line using the buttons in the upper left
#
# You're gonna wanna read this:
# https://code.visualstudio.com/docs/python/debugging
# Some of the pictures are a little dated

# Functions -------------------------------------------------------------------

# Pass in the file name to open:
def highlight_terms(file_name):
    """ Highlight some search terms in a given file.

        Blah blah blah example docstring. This is where most documentation
        for a function goes. If you mouse over the name of the function
        you'll see this description come up.

        Parameters:
        file_name (Path): A Path object to the given file
    """
    # Key words to find
    keywords = ['must', 'shall', 'will']

    # Use the file they input.
    the_doc = docx.Document(file_name)
    for pg in the_doc.paragraphs:
        for key in keywords:
            if key in pg.text:
                # inline = pg.runs
                for run in pg.runs:
                    x = run.text.split(key)
                    run.clear()
                    for i in range(len(x)-1):
                        run.add_text(x[i])
                        run.add_text(key)
                        # Line below causes an error
                        # run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
    the_doc.save('t2.docx')
    return 1

    # Here's a helpful SO I found on highlighting
    # https://stackoverflow.com/questions/46116389/highlight-text-using-python-docx



# Main ------------------------------------------------------------------------


if __name__ == '__main__':
    # How can we pass in the file name to here?
    # highlight_terms(the_file)
    
    # * Here's how we can take command line arguments
    # From the command line/terminal, you can run this script with
    # python highlight_search_terms.py path/to/file/here
    #
    # Two important libraries, sys and pathlib
    # 1.    pathlib provides a pretty easy to use way
    #       of guaranteeing that a given path to a file
    #       will work across all of the operating systems we use.
    #       Admittedly, not super useful in this script. We really don't use
    #       any of its features here, but it's important to know. I use it in sam_scraper.py and
    #       pdf_to_docx.py
    # 2.    sys (and "os") is a library with a bunch of random system stuff
    #       the most important things are sys.argv and sys.exit
    #       sys.argv is a list all arguments passed in on the command line
    # NOTE: sys.argv[0] is always the name of the script, so
    # "highlight_search_terms.py"

    if len(sys.argv) < 2:
        # If the only argument passed in was the script name
        # i.e. you gave no arguments
        print('Did you forget to pass in arguments?')
        sys.exit(1)
    else:
        # Turn the argument (type: str) into a Path object
        docx_path = Path(sys.argv[1])

        highlight_terms(docx_path)